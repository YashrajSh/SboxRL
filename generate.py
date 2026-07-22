"""
Futures First Intern OA Handbook Generator
Quant Intern - International Markets | SIT Pune
Run: python generate_handbook.py
Output: futures_first_handbook.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

                                                                                
for section in doc.sections:
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)

                                                                                
def set_font(run, size=11, bold=False, italic=False):
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    set_font(run, 14, bold=True)
                   
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, 12, bold=True)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, 11, bold=True, italic=True)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_font(run, 10)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, 10)
    return p

def formula(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, 10, bold=True)
    return p

def example(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(label + " ")
    set_font(r1, 10, bold=True)
    r2 = p.add_run(text)
    set_font(r2, 10)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    run = p.add_run("NOTE: " + text)
    set_font(run, 9, italic=True)
    return p

def page_break():
    doc.add_page_break()

def simple_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        set_font(run, 10, bold=True)
    for row_data in rows:
        row_cells = t.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            run = row_cells[i].paragraphs[0].runs[0]
            set_font(run, 9)
    doc.add_paragraph()

                                                                                
             
                                                                                
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run("FUTURES FIRST INTERN OA")
set_font(r, 22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Complete Aptitude & CS Handbook")
set_font(r, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Quant Intern  International Markets | SIT Pune")
set_font(r, 12, italic=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Covers: Mental Math | P&L | Ratio | Probability | DI | Number Systems\n"
              "Seating | Blood Relations | Coding-Decoding | Series | CS MCQs\n"
              "Python/NumPy/Pandas | 10 Mock Tests | Last-Hour Revision Sheet")
set_font(r, 10)

page_break()

                                                                                
                                    
                                                                                
h1("Section 1: Mental Math Shortcuts")

h2("1.1 Multiplication Tricks")

h3("Squaring numbers ending in 5")
formula("n5 = [n(n+1)] followed by 25")
example("Ex:", "75  78=56  5625.  95  910=90  9025")

h3("Multiply by 11")
formula("AB  11 = A (A+B) B  (carry if A+B >= 10)")
example("Ex:", "47  11  4 (4+7) 7 = 4(11)7  517.  63  11 = 693")

h3("Multiply any number by 9")
formula("n  9 = n  10  n")
example("Ex:", "47  9 = 470  47 = 423")

h3("Multiply two numbers near 100")
formula("(100a)(100b) = 100(100ab) + ab")
example("Ex:", "97  94  100(10036)+36 = 10091+18 = 9118")

h3("Cross-multiplication for 2-digit  2-digit")
formula("(10a+b)(10c+d) = 100ac + 10(ad+bc) + bd")
example("Ex:", "34  27  1006 + 10(8+21) + 28 = 600+290+28 = 918")

h3("Doubling and halving")
formula("a  b = (2a)  (b/2)   keep halving until easy")
example("Ex:", "125  48 = 12548  25024  50012  10006 = 6000")

h2("1.2 Division & Percentage Tricks")

h3("Key fraction-to-percent equivalents (MUST memorise)")
simple_table(
    ["Fraction","Percent","Fraction","Percent"],
    [
        ["1/8","12.5%","1/6","16.67%"],
        ["1/7","14.28%","1/9","11.11%"],
        ["1/11","9.09%","1/12","8.33%"],
        ["2/7","28.57%","3/7","42.86%"],
        ["3/8","37.5%","5/8","62.5%"],
        ["5/6","83.33%","7/8","87.5%"],
    ]
)

h3("% of a number  swap trick")
formula("X% of Y = Y% of X")
example("Ex:", "8% of 75 = 75% of 8 = 6  (easier to compute)")

h3("Successive percentage change")
formula("Net = a + b + ab/100   (a, b with signs)")
example("Ex:", "20% up then 20% down  20 + (20) + (2020)/100 = 0  4 = 4%")

h2("1.3 Squares, Cubes & Powers  Quick Reference")
simple_table(
    ["n","n","n","n","n","n"],
    [
        [1,1,1,11,121,1331],
        [2,4,8,12,144,1728],
        [3,9,27,13,169,2197],
        [4,16,64,14,196,2744],
        [5,25,125,15,225,3375],
        [6,36,216,16,256,4096],
        [7,49,343,17,289,4913],
        [8,64,512,18,324,5832],
        [9,81,729,19,361,6859],
        [10,100,1000,20,400,8000],
    ]
)

h2("1.4 Vedic Math  Base Multiplication")
formula("Numbers near base B: (B+a)(B+b) = B(B+a+b) + ab")
example("Ex (base 100):", "103  107 = 100(103+7) + 37 = 11000+21 = 11021")
example("Ex (base 10):", "8  7 = 10(8+710) + (810)(710) = 105+6 = 56")

h2("1.5 Approximation & Estimation")
bullet("Round to nearest 5 or power of 10, compute, adjust.")
bullet("For ratios: simplify numerator & denominator together.")
bullet("For large multiplications: log10 estimate  10^(log a + log b).")
bullet("Always check: units digit of answer (eliminates wrong MCQ options fast).")

page_break()

                                                                                
                            
                                                                                
h1("Section 2: Profit & Loss  Master Shortcuts")

h2("2.1 Core Formulas")
formula("SP = CP  (1  %/100)      [+ for profit,  for loss]")
formula("Profit% = (SPCP)/CP  100")
formula("Loss%   = (CPSP)/CP  100")
formula("CP given SP & Profit%:  CP = SP  100/(100+P%)")
formula("CP given SP & Loss%:    CP = SP  100/(100L%)")

h2("2.2 Discount & Marked Price")
formula("SP = MP  (1  d/100)")
formula("Profit% = [(MP/CP)(1d/100)  1]  100")
formula("For no profit no loss:  d = (MPCP)/MP  100")

h2("2.3 Dishonest Dealer Shortcut")
formula("Profit% = (True weight  False weight)/False weight  100")
example("Ex:", "Uses 900g weight instead of 1kg  Profit = 100/900  100 = 11.11%")

h2("2.4 Successive Discount")
formula("Equivalent single discount = a + b  ab/100")
example("Ex:", "30% + 20% = 30+206 = 44% (NOT 50%)")

h2("2.5 Solved Examples (25 sec)")
example("Q1:","CP=80, SP=100  Profit% = 20/80100 = 25%")
example("Q2:","SP=90, Loss=10%  CP = 90100/90 = 100")
example("Q3:","MP=200, Discount=15%  SP = 2000.85 = 170")
example("Q4:","Bought 10 for 1, sold 8 for 1  CP/unit=0.1, SP/unit=0.125  Profit=25%")
example("Q5:","Two articles at same SP. One 20%P, one 20%L  Net Loss% = (20)/100 = 4%")

note("Twin-article trap: same SP, one profit X% one loss X%  ALWAYS a net LOSS of X/100 %")

page_break()

                                                                                
                                           
                                                                                
h1("Section 3: Ratio, Proportion & Mixtures")

h2("3.1 Ratio Shortcuts")
formula("If a:b = m:n and b:c = p:q then a:b:c = mp : np : nq")
example("Ex:", "a:b=2:3, b:c=4:5  a:b:c = 8:12:15")
formula("Componendo-Dividendo: if a/b=c/d then (a+b)/(ab) = (c+d)/(cd)")

h2("3.2 Alligation Rule")
formula("Cheaper : Costlier = (MeanCheaper) : (CostlierMean)")
body("Draw the X diagram: diagonals give the ratio of quantities to mix.")
example("Ex:", "Milk at 15/L and 20/L mixed to get 17/L  ratio = (2017):(1715) = 3:2")

h2("3.3 Mixture & Replacement")
formula("After n replacements: Pure/Total = (1  r/V)^n")
example("Ex:", "20L vessel, remove 4L milk add water each time, after 3 rounds:")
example("","Milk fraction = (14/20) = (0.8) = 0.512  10.24L milk")

h2("3.4 Partnership")
formula("Profit share  Capital  Time")
example("Ex:", "A: 6000 for 6m, B: 4000 for 9m  ratio = 36000:36000 = 1:1")

page_break()

                                                                                
                          
                                                                                
h1("Section 4: Probability Tricks")

h2("4.1 Core Formulas")
formula("P(A) = Favourable / Total outcomes")
formula("P(AB) = P(A) + P(B)  P(AB)")
formula("P(A|B) = P(AB) / P(B)")
formula("Bayes: P(A|B) = P(B|A)P(A) / P(B)")

h2("4.2 Card Problems  Quick Reference")
body("Deck: 52 cards, 4 suits (13 each), 12 face cards, 4 aces")
simple_table(
    ["Event","Favourable","P"],
    [
        ["Ace",4,"4/52 = 1/13"],
        ["King",4,"1/13"],
        ["Face card",12,"3/13"],
        ["Red card",26,"1/2"],
        ["Red face card",6,"6/52 = 3/26"],
        ["Not a spade",39,"3/4"],
    ]
)

h2("4.3 Dice Problems")
body("Two dice: 36 total outcomes")
simple_table(
    ["Sum","Ways","P"],
    [
        [2,1,"1/36"],
        [3,2,"1/18"],
        [4,3,"1/12"],
        [5,4,"1/9"],
        [6,5,"5/36"],
        [7,6,"1/6"],
        [8,5,"5/36"],
    ]
)

h2("4.4 Complementary Counting (saves time)")
formula("P(at least one) = 1  P(none)")
example("Ex:", "P(at least one head in 3 tosses) = 1  (1/2) = 7/8")

h2("4.5 Permutation & Combination Fast Track")
formula("nPr = n!/(nr)!    nCr = n!/[r!(nr)!]")
formula("nCr = nC(nr)  [symmetry  use smaller r]")
formula("nC0=nCn=1, nC1=n")
example("Circular perm:","(n1)! arrangements")
example("With identical items:","n! / (p! q! r! ...) for p,q,r identical items")

h2("4.6 Solved Probability Examples")
example("Q1:","P(both red from 5R,3B without replacement) = 5/8  4/7 = 20/56 = 5/14")
example("Q2:","P(sum 8 with two dice) = 5/36")
example("Q3:","4 people, P(all birthdays different) = 365364363362/365  0.9836")
example("Q4 (Futures style):","P(stock up 3 days in a row, p=0.6) = 0.6 = 0.216")

page_break()

                                                                                
                                  
                                                                                
h1("Section 5: Data Interpretation Hacks")

h2("5.1 Speed Strategy")
bullet("Step 1: Read question FIRST, then scan table/chart for relevant rows only.")
bullet("Step 2: Never compute exact values if approximation separates options.")
bullet("Step 3: Use percentage change formula mentally  avoid long division.")
bullet("Step 4: For bar charts, read tallest/shortest before absolute values.")

h2("5.2 Percentage Change  2-Second Method")
formula("% Change = (Difference / Base)  100")
body("Trick: Express diff as fraction of base, convert using memorised fractions table.")
example("Ex:", "From 240 to 300  diff=60, 60/240 = 1/4 = 25%")

h2("5.3 CAGR Approximation")
formula("CAGR  (Final/Initial)^(1/n)  1")
body("Quick estimate: if value doubles in n years, CAGR  70/n  (Rule of 70)")
example("Ex:", "Doubles in 5 years  CAGR  14%")

h2("5.4 Pie Chart Tricks")
formula("Value = (Degree/360)  Total  OR  (Percent/100)  Total")
note("If two sectors differ by d degrees, their value difference = d/360  Total")

h2("5.5 Sample DI Set  Table")
body("Company revenues ( Crore):")
simple_table(
    ["Company","2022","2023","2024"],
    [
        ["Alpha",500,600,750],
        ["Beta",300,360,432],
        ["Gamma",200,220,242],
        ["Delta",400,480,576],
    ]
)
example("Q1 (Growth %):", "Alpha 2223: 600500=100, 100/500=20%. Beta: 360300=60, 60/300=20%.")
example("Q2 (CAGR Alpha 2224):","750/500 = 1.5. CAGR = 1.5^0.5  1  22.5%")
example("Q3 (Ratio 2024 Alpha:Gamma):","750:242  375:121  3.1:1")

page_break()

                                                                                
                             
                                                                                
h1("Section 6: Number Systems  HCF, LCM, Divisibility, Remainders")

h2("6.1 HCF & LCM")
formula("HCF  LCM = Product of two numbers (ONLY for two numbers)")
formula("HCF of fractions = HCF(numerators) / LCM(denominators)")
formula("LCM of fractions = LCM(numerators) / HCF(denominators)")
note("For 3+ numbers, HCFLCM  product. Use prime factorisation.")

h2("6.2 Divisibility Rules")
simple_table(
    ["Divisor","Rule"],
    [
        [2,"Last digit even"],
        [3,"Sum of digits divisible by 3"],
        [4,"Last 2 digits divisible by 4"],
        [6,"Divisible by both 2 and 3"],
        [7,"Double last digit, subtract from rest; repeat"],
        [8,"Last 3 digits divisible by 8"],
        [9,"Sum of digits divisible by 9"],
        [11,"Alternating sum of digits divisible by 11"],
        [12,"Divisible by both 3 and 4"],
        [25,"Last 2 digits = 00, 25, 50, or 75"],
    ]
)

h2("6.3 Remainder Theorems")
h3("Fermat's Little Theorem")
formula("a^(p1)  1 (mod p)  where p is prime and gcd(a,p)=1")
example("Ex:", "7^100 mod 5  72(mod5), 2^41(mod5), 100=425  2^100=(2^4)^25  1(mod5)")

h3("Cyclicity of Units Digits")
simple_table(
    ["Base (units)","Cycle length","Pattern"],
    [
        [0,"1","0"],
        [1,"1","1"],
        [2,"4","2,4,8,6"],
        [3,"4","3,9,7,1"],
        [4,"2","4,6"],
        [5,"1","5"],
        [6,"1","6"],
        [7,"4","7,9,3,1"],
        [8,"4","8,4,2,6"],
        [9,"2","9,1"],
    ]
)
example("Ex:", "7^355  355 mod 4 = 3  3rd in cycle (7,9,3,1) = 3. Units digit = 3")

h2("6.4 Remainders  Chinese Remainder Theorem (simple)")
formula("Find x: x mod a = r1, x mod b = r2   solve by listing multiples")
example("Ex:", "x mod 3=2, x mod 5=3  x=8,23,38...  smallest=8")

h2("6.5 Number of Factors")
formula("n = p1^a  p2^b  p3^c  factors = (a+1)(b+1)(c+1)")
example("Ex:", "360 = 235  factors = 432 = 24")
formula("Sum of factors = (p1^(a+1)1)/(p11)  (p2^(b+1)1)/(p21)  ...")

h2("6.6 LCM Word Problems  Template")
formula("When does cycle repeat?  LCM of individual periods")
example("Ex:", "Bells ring every 12, 15, 20 min  LCM(12,15,20) = 60 min")

page_break()

                                                                                
                                
                                                                                
h1("Section 7: Time, Speed, Distance & Work")

h2("7.1 Speed Shortcuts")
formula("Average speed (two equal distances) = 2xy/(x+y)  [harmonic mean]")
example("Ex:", "60 km/h one way, 40 km/h return  avg = 26040/100 = 48 km/h")
formula("Relative speed: same dir = |v1v2|, opposite dir = v1+v2")
formula("Train crossing pole: time = length/speed. Crossing platform: (L+P)/speed")

h2("7.2 Work Shortcuts")
formula("A does in a days, B in b days  together in ab/(a+b) days")
formula("If A is twice as fast as B  ratio of time = 1:2")
example("Ex:", "A=10d, B=15d  together = 1015/25 = 6 days")
formula("Pipes: filling rates add, draining rates subtract")
example("Ex:", "Fill in 6h, drain in 12h  net = 1/6  1/12 = 1/12  12 hours")

h2("7.3 Races")
formula("A beats B by d metres in race of L  B covers (Ld) when A covers L")
formula("A's speed / B's speed = L / (Ld)")
example("Ex:", "In 100m race, A beats B by 10m  ratio=100:90=10:9")

page_break()

                                                                                
                                
                                                                                
h1("Section 8: Logical Reasoning")

h2("8.1 Seating Arrangement  Templates")

h3("Linear Arrangement  Steps")
bullet("Step 1: Fix the most constrained person (most clues about them).")
bullet("Step 2: Use definite clues first (exact positions).")
bullet("Step 3: Use relative clues (A is 2nd to the left of B).")
bullet("Step 4: Plug in remaining by elimination.")
body("Key patterns:")
bullet("'Immediate neighbours' = exactly 1 apart")
bullet("'Not adjacent' = at least 2 apart")
bullet("'Facing each other' in circular = directly opposite")

h3("Circular Arrangement  Steps")
bullet("Fix one person as reference (say Person A at top).")
bullet("Assign others relative to A using clues.")
bullet("'To the left' = anticlockwise in standard problems (confirm from clue).")
bullet("n people in circle: (n1)! distinct arrangements")

h3("Double-Row Arrangement")
bullet("Draw two rows, establish who faces whom.")
bullet("'Faces A' means in opposite row, same column.")
bullet("'To the left of person facing B' = right of B's row position.")

h2("8.2 Blood Relations  Master Chart")
body("Use M=Male, F=Female, U=Unknown. Draw the family tree.")
simple_table(
    ["Relation said","Meaning"],
    [
        ["Father's brother","Uncle (Paternal)"],
        ["Mother's brother","Maternal Uncle"],
        ["Father's sister's husband","Uncle"],
        ["Son's wife","Daughter-in-law"],
        ["Wife's brother","Brother-in-law"],
        ["Father's father","Paternal Grandfather"],
        ["Brother's son","Nephew"],
        ["Sister's daughter","Niece"],
        ["Husband's/wife's father","Father-in-law"],
    ]
)
h3("Chain Rule for coded relations")
formula("A's B = X.  X's C = Y.  Find A's relation to Y.")
body("Draw: A  (via B)  X  (via C)  Y. Trace path using gender.")
example("Ex:", "A's father's sister's son = A's cousin (paternal)")

h2("8.3 Coding-Decoding Patterns")
h3("Letter shift coding")
formula("Each letter shifted by fixed n: AD means shift=3. Decode by reversing shift.")
example("Ex:", "COLD  FROG = each letter +3")

h3("Position reversal")
body("COME  XLNV = reverse alphabet (AZ, BY, CX...)")
formula("Coded letter = 27  position of original letter")

h3("Number coding")
body("Letters assigned numbers: A=1,B=2...Z=26 or A=26...Z=1")
bullet("Look at first few letters and their codes to determine the pattern.")
bullet("Check if it's position, position+constant, or a completely different mapping.")

h3("Symbol/mixed coding")
bullet("Find the common element between two sentences with one common word.")
bullet("The code for the common word = common symbol in both coded sentences.")

h2("8.4 Series Recognition Tricks")
h3("Number Series")
bullet("Difference series: find 1st diff, 2nd diff, 3rd diff. Look for AP/GP in differences.")
bullet("Product series: each term = previous  some factor (check ratios).")
bullet("Square/Cube series: check if terms are n, n, n+k, etc.")
bullet("Twin series: alternate terms form two separate series.")

simple_table(
    ["Pattern","Example","Rule"],
    [
        ["Perfect squares","1,4,9,16,25","n"],
        ["Fibonacci","1,1,2,3,5,8,13","prev+prev-prev"],
        ["Powers of 2","2,4,8,16,32","2"],
        ["2+1","3,7,15,31,63","2+1"],
        ["n(n+1)","2,6,12,20,30","n(n+1)"],
        ["Prime","2,3,5,7,11,13","primes"],
    ]
)

h3("Letter Series")
bullet("Count positions skipped between consecutive letters.")
bullet("Look for +1,+2,+3 or alternating patterns.")
example("Ex:", "B, E, H, K  +3 each time  next = N")

h2("8.5 Mirror & Image Reasoning")
h3("Clock reflection (mirror on right)")
formula("Mirror time = 11:60  actual time")
example("Ex:", "Clock shows 8:45  Mirror = 11:60  8:45 = 3:15")
h3("Letter/Image mirror")
bullet("Vertical mirror: left-right flip. Horizontal mirror: top-bottom flip.")
bullet("For alphabets in mirror: B (reverse B), E reversed E, etc.")
bullet("Number trick: only 1,8,0 look same in vertical mirror.")

page_break()

                                                                                
                                 
                                                                                
h1("Section 9: Common Quantitative Traps")

simple_table(
    ["Trap","Wrong approach","Correct approach"],
    [
        ["Average speed","(v1+v2)/2","2v1v2/(v1+v2)"],
        ["Two articles same SP, X% profit and X% loss","0% net change","Net LOSS = X/100 %"],
        ["Successive % change","add %","a+b+ab/100"],
        ["% more/less confusion","use same base","specify base clearly"],
        ["Simple vs Compound interest","same formula","CI formula has (1+r/100)^n"],
        ["Discount on Discount","add them","Equivalent = a+bab/100"],
        ["Ratio: if a:b=2:3, find a+b","2+3=5 parts","use multiplier, check total"],
        ["Probability of at least 1","direct count","1  P(none)"],
        ["Circular permutation","n!","(n1)!"],
        ["HCF  LCM = product","works for any n numbers","only for exactly 2 numbers"],
    ]
)

page_break()

                                                                                
                       
                                                                                
h1("Section 10: CS MCQs  OOP, DBMS, OS, CN, Java, SQL")

h2("10.1 OOP Concepts")
simple_table(
    ["Concept","Definition","Key point"],
    [
        ["Encapsulation","Bundling data+methods","Access via public methods only"],
        ["Abstraction","Hide implementation detail","Abstract class/Interface"],
        ["Inheritance","Child inherits parent","IS-A relationship"],
        ["Polymorphism","One interface, many forms","Overloading(compile) / Overriding(runtime)"],
        ["Constructor","Init method, same name as class","No return type"],
        ["Destructor","Cleanup method","~ClassName() in C++"],
    ]
)
bullet("Method overloading: same name, different parameters (compile-time).")
bullet("Method overriding: same signature, child class redefines (runtime).")
bullet("Abstract class: cannot instantiate; can have concrete methods.")
bullet("Interface: all methods abstract (Java 7); default methods allowed (Java 8+).")
bullet("Multiple inheritance: C++ supports; Java does NOT (use interfaces).")

h2("10.2 DBMS")
h3("Normalisation")
simple_table(
    ["Normal Form","Condition"],
    [
        ["1NF","Atomic values, no repeating groups"],
        ["2NF","1NF + no partial dependency (on composite key)"],
        ["3NF","2NF + no transitive dependency"],
        ["BCNF","For every FD XY, X must be a superkey"],
    ]
)
h3("SQL Keywords  Fast Reference")
simple_table(
    ["Clause","Use"],
    [
        ["SELECT DISTINCT","Remove duplicates"],
        ["GROUP BY","Aggregate rows"],
        ["HAVING","Filter after GROUP BY (not WHERE)"],
        ["INNER JOIN","Matching rows only"],
        ["LEFT JOIN","All left + matching right"],
        ["FULL OUTER JOIN","All rows from both"],
        ["UNION","Combine, remove duplicates"],
        ["UNION ALL","Combine, keep duplicates"],
        ["EXISTS","Subquery returns rows"],
        ["COALESCE(a,b)","Return first non-NULL"],
    ]
)
body("ACID properties: Atomicity, Consistency, Isolation, Durability")
body("Indexing: B+ tree default. Speeds SELECT, slows INSERT/UPDATE.")

h2("10.3 Operating Systems")
simple_table(
    ["Topic","Key Facts"],
    [
        ["Process vs Thread","Process: own memory. Thread: shared memory, lighter."],
        ["Deadlock conditions","Mutual exclusion, Hold & wait, No preemption, Circular wait (ALL 4 needed)"],
        ["Page replacement","FIFO, LRU (least recently used), Optimal (Belady's)"],
        ["Scheduling","FCFS, SJF (non-preemptive), SRTF (preemptive), Round Robin, Priority"],
        ["Thrashing","Excessive paging, process spends more time paging than executing"],
        ["Semaphore","wait()/signal(); mutex=binary semaphore"],
        ["Virtual memory","Allows programs larger than RAM; uses paging/segmentation"],
    ]
)

h2("10.4 Computer Networks")
simple_table(
    ["Topic","Key Facts"],
    [
        ["OSI Layers","Physical, Data Link, Network, Transport, Session, Presentation, Application"],
        ["TCP vs UDP","TCP: reliable, ordered, connection-oriented. UDP: fast, no guarantee."],
        ["HTTP vs HTTPS","HTTPS = HTTP + TLS/SSL (port 443 vs 80)"],
        ["IP addressing","IPv4: 32-bit. IPv6: 128-bit. Subnet mask separates network/host"],
        ["DNS","Domain  IP translation. Recursive vs iterative queries."],
        ["3-way handshake","SYN  SYN-ACK  ACK (TCP connection setup)"],
        ["Routing protocols","RIP (distance vector), OSPF (link state), BGP (path vector)"],
    ]
)

h2("10.5 Java Quick Facts")
bullet("Java is platform-independent: compiled to bytecode  JVM runs it.")
bullet("Primitive types: byte, short, int, long, float, double, char, boolean.")
bullet("String is immutable; StringBuilder/StringBuffer are mutable.")
bullet("'==' compares references; '.equals()' compares values.")
bullet("ArrayList: dynamic array (random access O(1)). LinkedList: O(n) access, O(1) insert at ends.")
bullet("HashMap: O(1) average get/put; allows one null key. TreeMap: sorted, O(log n).")
bullet("final: variable=constant, method=no override, class=no extend.")
bullet("static: belongs to class, not instance.")
bullet("Checked exceptions must be caught/declared. Unchecked (RuntimeException) need not be.")
bullet("Garbage collection: automatic; System.gc() is a hint, not a guarantee.")
bullet("Autoboxing: int  Integer automatic conversion.")

h2("10.6 SQL Queries  Practice")
body("Find 2nd highest salary:")
formula("SELECT MAX(salary) FROM emp WHERE salary < (SELECT MAX(salary) FROM emp);")
body("Count employees per department:")
formula("SELECT dept, COUNT(*) FROM emp GROUP BY dept;")
body("Employees earning more than dept average:")
formula("SELECT * FROM emp e WHERE salary > (SELECT AVG(salary) FROM emp WHERE dept=e.dept);")
body("Delete duplicates keeping one:")
formula("DELETE FROM emp WHERE id NOT IN (SELECT MIN(id) FROM emp GROUP BY name, salary);")

page_break()

                                                                                
                                            
                                                                                
h1("Section 11: Time Complexity Shortcuts")

h2("11.1 Big-O Quick Reference")
simple_table(
    ["Complexity","Name","Example"],
    [
        ["O(1)","Constant","Array index, HashMap get"],
        ["O(log n)","Logarithmic","Binary search, balanced BST"],
        ["O(n)","Linear","Linear search, single loop"],
        ["O(n log n)","Log-linear","Merge sort, Heap sort, Quick sort (avg)"],
        ["O(n)","Quadratic","Bubble/Selection/Insertion sort, nested loops"],
        ["O(n)","Cubic","Floyd-Warshall"],
        ["O(2^n)","Exponential","Fibonacci (naive), subsets"],
        ["O(n!)","Factorial","Permutations, TSP brute force"],
    ]
)

h2("11.2 Sorting Algorithm Summary")
simple_table(
    ["Algorithm","Best","Average","Worst","Space","Stable?"],
    [
        ["Bubble Sort","O(n)","O(n)","O(n)","O(1)","Yes"],
        ["Selection Sort","O(n)","O(n)","O(n)","O(1)","No"],
        ["Insertion Sort","O(n)","O(n)","O(n)","O(1)","Yes"],
        ["Merge Sort","O(n log n)","O(n log n)","O(n log n)","O(n)","Yes"],
        ["Quick Sort","O(n log n)","O(n log n)","O(n)","O(log n)","No"],
        ["Heap Sort","O(n log n)","O(n log n)","O(n log n)","O(1)","No"],
        ["Counting Sort","O(n+k)","O(n+k)","O(n+k)","O(k)","Yes"],
    ]
)

h2("11.3 Data Structure Operations")
simple_table(
    ["Structure","Access","Search","Insert","Delete"],
    [
        ["Array","O(1)","O(n)","O(n)","O(n)"],
        ["Linked List","O(n)","O(n)","O(1)","O(1)"],
        ["Stack","O(n)","O(n)","O(1)","O(1)"],
        ["Queue","O(n)","O(n)","O(1)","O(1)"],
        ["Hash Table","N/A","O(1) avg","O(1) avg","O(1) avg"],
        ["BST (balanced)","O(log n)","O(log n)","O(log n)","O(log n)"],
        ["Heap","N/A","O(n)","O(log n)","O(log n)"],
    ]
)

h2("11.4 Graph Algorithm Complexities")
simple_table(
    ["Algorithm","Time","Space","Use"],
    [
        ["BFS","O(V+E)","O(V)","Shortest path (unweighted)"],
        ["DFS","O(V+E)","O(V)","Cycle detection, topological sort"],
        ["Dijkstra","O((V+E)log V)","O(V)","Shortest path (positive weights)"],
        ["Bellman-Ford","O(VE)","O(V)","Negative weights OK"],
        ["Floyd-Warshall","O(V)","O(V)","All-pairs shortest path"],
        ["Kruskal","O(E log E)","O(V)","MST"],
        ["Prim","O(E log V)","O(V)","MST"],
    ]
)

page_break()

                                                                                
                                              
                                                                                
h1("Section 12: Python / NumPy / Pandas Revision")

h2("12.1 Python Data Structures  Critical Facts")
simple_table(
    ["Type","Mutable","Ordered","Duplicate","Syntax"],
    [
        ["list","Yes","Yes","Yes","[1,2,3]"],
        ["tuple","No","Yes","Yes","(1,2,3)"],
        ["set","Yes","No","No","{1,2,3}"],
        ["dict","Yes","Yes(3.7+)","Keys:No","{k:v}"],
        ["frozenset","No","No","No","frozenset({1,2})"],
    ]
)

h2("12.2 List Comprehension & Lambda")
formula("[x**2 for x in range(10) if x%2==0]  # even squares")
formula("sorted(lst, key=lambda x: x[1])  # sort by 2nd element")
formula("map(func, lst)  # lazy; filter(func, lst)")

h2("12.3 NumPy  Key Operations")
body("Import: import numpy as np")
simple_table(
    ["Operation","Code"],
    [
        ["Create array","np.array([1,2,3])"],
        ["Zeros/Ones","np.zeros((3,4)), np.ones((2,2))"],
        ["Range","np.arange(0,10,2), np.linspace(0,1,100)"],
        ["Shape/Reshape","a.shape, a.reshape(2,3)"],
        ["Element-wise ops","a+b, a*b, a**2 (broadcasting)"],
        ["Dot product","np.dot(a,b) or a@b"],
        ["Transpose","a.T"],
        ["Axis operations","np.sum(a,axis=0), np.mean(a,axis=1)"],
        ["Boolean indexing","a[a>5]"],
        ["Stack","np.vstack([a,b]), np.hstack([a,b])"],
        ["Random","np.random.rand(3,3), np.random.randn(100)"],
        ["Statistics","np.mean, np.std, np.var, np.median, np.percentile"],
    ]
)

h2("12.4 Pandas  Key Operations")
body("Import: import pandas as pd")
simple_table(
    ["Operation","Code"],
    [
        ["Read CSV","pd.read_csv('file.csv')"],
        ["Head/Tail","df.head(), df.tail(10)"],
        ["Info/Describe","df.info(), df.describe()"],
        ["Select column","df['col'] or df[['c1','c2']]"],
        ["Filter rows","df[df['age']>30]"],
        ["GroupBy","df.groupby('col').mean()"],
        ["Merge","pd.merge(df1,df2,on='key',how='inner')"],
        ["Pivot table","df.pivot_table(values='v',index='r',columns='c',aggfunc='sum')"],
        ["Handle nulls","df.dropna(), df.fillna(0), df.isna().sum()"],
        ["Apply function","df['col'].apply(lambda x: x*2)"],
        ["Sort","df.sort_values('col',ascending=False)"],
        ["Reset index","df.reset_index(drop=True)"],
        ["String ops","df['col'].str.upper(), .str.contains('pat')"],
        ["Date ops","pd.to_datetime(df['date']), df['date'].dt.year"],
    ]
)

h2("12.5 Python MCQ Traps")
bullet("'is' checks identity (same object); '==' checks equality.")
bullet("Mutable default argument: def f(x=[]) is a classic bug  list persists across calls.")
bullet("Global scope: must use 'global x' inside function to modify global x.")
bullet("Shallow vs deep copy: list.copy() / copy.copy() = shallow; copy.deepcopy() = deep.")
bullet("Generator vs list: generator is lazy (yield), saves memory.")
bullet("*args = tuple of positional args; **kwargs = dict of keyword args.")
bullet("__str__ = user-facing string; __repr__ = developer-facing representation.")
bullet("Slicing: lst[a:b:c]  step c; negative step reverses.")
bullet("'in' operator: O(n) for list, O(1) average for set/dict.")

page_break()

                                                                                
                                                               
                                                                                
h1("Section 13: Quant Finance & Trading Concepts")

h2("13.1 Basic Finance MCQ Topics")
simple_table(
    ["Topic","Key Fact"],
    [
        ["Futures","Obligation to buy/sell at future date. Exchange-traded."],
        ["Options","Right, not obligation. Call=buy, Put=sell."],
        ["Arbitrage","Risk-free profit from price discrepancy."],
        ["Bid-Ask spread","Ask > Bid. Spread = Ask  Bid."],
        ["Long/Short","Long = bought = profit if price rises. Short = sold = profit if price falls."],
        ["Leverage","Control large position with small capital. Amplifies gains AND losses."],
        ["Mark to Market","Daily settlement of futures based on end-of-day price."],
        ["Basis","Spot price  Futures price (or vice versa depending on convention)"],
        ["Contango","Futures > Spot (normal curve, storage costs)"],
        ["Backwardation","Futures < Spot (shortage, convenience yield dominates)"],
        ["P&L for Long future","(Exit price  Entry price)  lot size"],
        ["Greeks (Option)","Delta=price sensitivity, Gamma=delta change, Theta=time decay, Vega=vol sensitivity"],
    ]
)

h2("13.2 Expected Value  Trader's Tool")
formula("E[X] =  p_i  x_i")
example("Ex:", "Bet $10: win $30 with p=0.4, lose $10 with p=0.6  E = 0.430 + 0.6(10) = 126 = $6 (take it)")
body("Futures First heavily tests EV reasoning and bet sizing.")

h2("13.3 Bayes & Conditional Probability in Trading")
formula("P(profit | signal) = P(signal | profit)  P(profit) / P(signal)")
example("Ex:", "Strategy profitable 70% of time. Signal fires 60% of days.")
example("","P(profit|signal) depends on P(signal|profit)  must be specified.")

h2("13.4 Mental Math for Trading")
bullet("PnL calculation: always in units  price move  lot size.")
bullet("Annualise: daily vol  sqrt(252). Monthly vol  sqrt(12).")
bullet("Sharpe Ratio = (Return  Risk-free) / Std Dev of returns.")
bullet("If coin is unfair (p heads  0.5): expected # tosses for first head = 1/p.")
bullet("Kelly Criterion: bet fraction = (bpq)/b  where b=odds, p=win prob, q=1p.")

page_break()

                                                                                
                             
                                                                                
h1("Section 14: 10 Futures First-Style Mock Tests")

def mock_test(num, questions):
    h2(f"Mock Test {num}")
    for i, (q, a) in enumerate(questions, 1):
        body(f"Q{i}. {q}")
        example("Ans:", a)
    doc.add_paragraph()

        
mock_test(1, [
    ("If 6 is 15% of X, find X.", "X = 6/0.15 = 40"),
    ("A train 120m long passes a pole in 6 seconds. Speed?", "120/6 = 20 m/s = 72 km/h"),
    ("What is 7^(50) mod 4?", "73(mod4). 3^2=91(mod4). 50=225  3^50=(3^2)^251(mod4). Ans: 1"),
    ("P(at least one 6 in 2 dice rolls)?", "1(5/6) = 125/36 = 11/36"),
    ("2 pipes fill tank in 6h and 9h. Drain pipe empties in 18h. All open  time to fill?", "Rate=1/6+1/91/18 = 3/18+2/181/18=4/18=2/9. Time=9/2=4.5 hours"),
    ("Find HCF of 72 and 120.", "72=23, 120=235. HCF=23=24"),
    ("X% of 80 = 20. X?", "20/80100 = 25"),
    ("Class: Avg marks 40, pass marks 50% of max=50. Avg of passed = 60, failed=20. Passed?", "Let pass=p, fail=f. 60p+20f=40(p+f)  20p=20f  p=f  50% passed."),
    ("What is the units digit of 3^123?", "Cycle of 3: 3,9,7,1 (period 4). 123 mod 4=3  3rd in cycle = 7"),
    ("Find LCM of 12, 18, 24.", "12=23, 18=23, 24=23. LCM=23=72"),
])

        
mock_test(2, [
    ("Ratio of ages A:B=3:4. After 8 years, ratio=5:6. Current ages?", "3x+8)/(4x+8)=5/6  18x+48=20x+40  2x=8  x=4. A=12, B=16"),
    ("A sells to B at 20% profit. B sells to C at 10% loss. If C pays 2160, what did A pay?", "C's price = A's cost  1.2  0.9 = 1.08cost. cost=2160/1.08=2000"),
    ("Find the odd one out: 2, 5, 10, 17, 26, 37, 50, 65", "Differences: 3,5,7,9,11,13,15 (odd). All fit n+1. Actually 3750=13, 5065=15. All correct. Odd one: None  trick Q. If forced: check 65=8+1=65 "),
    ("P(sum>9 with two dice)?", "Favourable: (4,6),(5,5),(5,6),(6,4),(6,5),(6,6)=6 ways. P=6/36=1/6"),
    ("Speed of boat in still water 15km/h. Stream 3km/h. Time for 54km downstream?", "Downstream=18km/h. Time=54/18=3 hours"),
    ("ABCD is coded as BCDA. What is MNOP coded as?", "Each letter shifts 1 position forward cyclically. MNOP  NOPM"),
    ("If 8 workers complete job in 15 days, how many workers complete it in 6 days?", "Total work=120 man-days. Workers=120/6=20"),
    ("Two numbers have HCF=12, LCM=144. One number is 36. Other?", "Other = HCFLCM/36 = 12144/36 = 48"),
    ("Compound interest on 8000 at 10% per annum for 2 years?", "8000(1.1)=80001.21=9680. CI=1680"),
    ("Which SQL clause filters groups? (a) WHERE (b) HAVING (c) GROUP BY (d) ORDER BY", "(b) HAVING  filters after GROUP BY"),
])

        
mock_test(3, [
    ("In what ratio water be mixed with milk to gain 20% by selling at cost price?", "Profit from adulteration: gain=20%1/5. Water:Milk = 1:5"),
    ("A can do work in 10 days, B in 15. They start together but A leaves after 3 days. Total days?", "Work done in 3d = 3(1/10+1/15)=31/6=1/2. Remaining done by B: (1/2)/(1/15)=7.5d. Total=10.5 days"),
    ("Expected value: 60% chance of gaining $50, 40% of losing $30. EV?", "E=0.650+0.4(30)=3012=$18"),
    ("Find the missing: 3, 6, 11, 18, 27, ?", "Diff: 3,5,7,9,11. Next diff=11 Next=27+11=38"),
    ("A person walks 5km N, 3km E, 5km S. Distance from start?", "N and S cancel. Net displacement=3km East"),
    ("N facing, turn 90 right, then 45 left. Now facing?", "Nright=E. Eleft 45=NE. Facing NE (North-East)"),
    ("Virtual memory allows programs larger than RAM. True/False?", "True  OS swaps pages between RAM and disk."),
    ("Deadlock needs all 4 conditions simultaneously. Remove one to prevent. Which is easiest?","Remove 'Hold and Wait': process must acquire all resources at once."),
    ("What does HAVING differ from WHERE?", "WHERE filters rows before aggregation. HAVING filters groups after GROUP BY."),
    ("Time complexity of binary search?", "O(log n)"),
])

        
mock_test(4, [
    ("Price increased 25%, then decreased 20%. Net change?", "25+(20)+25(20)/100=55=0%. No net change."),
    ("Bag: 4R, 3G, 2B balls. P(2 balls same colour without replacement)?", "Total C(9,2)=36. Same: C(4,2)+C(3,2)+C(2,2)=6+3+1=10. P=10/36=5/18"),
    ("Average of 5 numbers is 27. If one number is excluded, avg becomes 25. Excluded number?","Sum=135. New sum=254=100. Excluded=35"),
    ("1 is the father of 2. 3 is the son of 2. 4 is the brother of 1. How is 4 related to 3?","4=brother of 1=grandfather of 3. 4 is great-uncle? No: 1 is father of 2, 2 is parent of 3. 4 is brother of 1  uncle of 2  great-uncle of 3. Ans: 4 is great-uncle of 3."),
    ("What is thrashing in OS?", "Process spends more time paging than executing; too many processes competing for limited frames."),
    ("Difference between TCP and UDP in one line?", "TCP: reliable, connection-oriented, ordered. UDP: fast, connectionless, no delivery guarantee."),
    ("Which normal form removes transitive dependency?", "3NF"),
    ("n=12! How many trailing zeros?", "Trailing zeros=min(powers of 2 and 5) in n!. Fives: floor(12/5)=2. Zeros=2"),
    ("COLD  WARM (in some code). BIRDS  ?", "CW(+4), OA(+3)... not constant. Try: reverse COLD=DLOC. Probably problem needs exact coding. Pattern: COLD's mirror in alphabet=XLOW. Needs more info."),
    ("Sharpe Ratio formula?", "Sharpe = (Portfolio Return  Risk-free Rate) / Std Dev of excess return"),
])

        
mock_test(5, [
    ("Sum of first n odd numbers?", "n (e.g., 1+3+5+7=16=4)"),
    ("In 6 years A will be twice as old as B was 6 years ago. Currently A=50. B?","A in 6 years=56. B 6 years ago=28. B now=34."),
    ("P(at least 2 heads in 3 fair coin tosses)?","P(exactly 2)+P(exactly 3)=C(3,2)/8+1/8=3/8+1/8=4/8=1/2"),
    ("Interface vs Abstract class in Java?","Interface: all abstract methods (pre Java 8), no state. Abstract class: can have concrete methods and fields. A class can implement multiple interfaces but extend only one abstract class."),
    ("What is a hash collision and how is it handled?","Two keys hash to same index. Handled by: chaining (linked list at each bucket) or open addressing (probe next slot)."),
    ("Time complexity of Merge Sort in all cases?","O(n log n) in all cases  best, average, worst."),
    ("Find inverse: A buys at SP of 16 articles for cost of 20. Profit%?","CP for 20 = SP for 16. CP per unit = SP16/20. Profit% = (2016)/16100=25%"),
    ("Two dice: P(product is even)?","P(product odd) = P(both odd)=(3/6)=1/4. P(even)=11/4=3/4"),
    ("What is a deadlock-free situation called?","Starvation-free AND deadlock-free: Liveness property. Deadlock prevention/avoidance ensures deadlock-free state."),
    ("Kelly Criterion: win prob=0.6, odds=1:1 (bet 1 to win 1). Optimal bet fraction?","f = (bpq)/b = (10.60.4)/1 = 0.2 = 20% of bankroll"),
])

        
mock_test(6, [
    ("A number when divided by 5 gives remainder 2, by 7 gives remainder 3. Smallest such number?","By 5: 2,7,12,17,22,27,32,37. By 7: 3,10,17,24,31,38. Common: 17"),
    ("Boat speed 12km/h still water, stream 4km/h. Time ratio upstream:downstream for same distance?","Upstream=8, Downstream=16. Ratio=16:8=2:1"),
    ("A series: 2,3,5,8,13,21. Next?","Fibonacci: 21+13=34"),
    ("What does ACID stand for in DBMS?","Atomicity, Consistency, Isolation, Durability"),
    ("In Python: x=[1,2,3]; y=x; y.append(4). What is x?","[1,2,3,4]  y=x is a reference copy, same object."),
    ("numpy: What does a.reshape(-1) do?","Flattens array to 1D. -1 infers the dimension."),
    ("Logical: All A are B. Some B are C. Therefore?","Some A may or may not be C. No definite conclusion about A and C."),
    ("Three persons A,B,C sit in circle. B sits to the right of A. C sits to the left of A. Who sits between B and C (not through A)?","A sits between B and C (through A). On the other side, B? C: only A can be between them clockwise or anticlockwise."),
    ("What is contango in futures markets?","Futures price > Spot price. Typical in markets with storage costs."),
    ("Quick sort worst case and when does it occur?","O(n) when pivot is always smallest/largest element (sorted array with first/last pivot)."),
])

        
mock_test(7, [
    ("Mirror time: clock shows 10:10. What does mirror show?","11:6010:10=1:50"),
    ("A can finish work in 12 days. After 4 days, B joins. Together finish in 4 more days. B alone?","A does 4/12=1/3 in 4 days. Remaining=2/3. In 4 days, A does 4/12=1/3. B does 2/31/3=1/3 in 4 days. B alone: 12 days."),
    ("Find the number of divisors of 180.","180=235. Divisors=(2+1)(2+1)(1+1)=18"),
    ("SQL: Find employees with salary higher than their manager's.","SELECT e.name FROM emp e JOIN emp m ON e.manager_id=m.id WHERE e.salary > m.salary"),
    ("P(king or queen from a deck of cards)?","8/52 = 2/13"),
    ("In Python, what is the output of print(type(1/2))?","<class 'float'>  Python 3 division always returns float"),
    ("What is the difference between == and is in Python?","== compares values. is compares identity (same memory address)."),
    ("What is the time complexity of inserting into a sorted linked list?","O(n) to find position, O(1) to insert. Total: O(n)"),
    ("Decode: In a code, MANGO=13,1,14,7,15. GRAPE=?","A=1,B=2...positional. G=7,R=18,A=1,P=16,E=5. GRAPE=7,18,1,16,5"),
    ("What is Vega in options?","Sensitivity of option price to changes in implied volatility of the underlying."),
])

        
mock_test(8, [
    ("Mixture has milk:water = 4:1. Add 10L water to make ratio 2:3. Original mixture?","Let orig milk=4x, water=x. (4x)/(x+10) = 2/3  12x=2x+20  10x=20  x=2. Orig total=10L"),
    ("CI on 10000 at 5% pa for 3 years?","10000(1.05)=100001.157625=11576.25. CI=1576.25"),
    ("Python: output of [i*2 for i in range(5) if i%2!=0]?","[2,6]  i=1,3 (odd), 2"),
    ("What is polymorphism? Give one Java example.","One interface, multiple implementations. Example: method overriding  Animal.speak() defined differently in Dog and Cat."),
    ("Find 2 numbers whose HCF=4, LCM=48, and one is 12.","Other = HCFLCM/12 = 448/12 = 16"),
    ("P(both cards red when 2 drawn from deck without replacement)?","26/52  25/51 = 650/2652 = 25/102"),
    ("Seating: 6 people in circle, 2 specific people must always sit together. Arrangements?","Treat pair as one unit: 5 units in circle = (51)!=24. Pair can arrange internally in 2!=2 ways. Total=48"),
    ("NumPy: difference between np.dot and np.matmul?","For 2D arrays they are same. For 3D+, matmul does batch matrix multiplication. dot does sum-product along last axis."),
    ("A stock goes up 10% then down 10%. Net change?","1.1  0.9 = 0.99. Net = 1%"),
    ("BGP stands for? Used for?","Border Gateway Protocol. Inter-AS (autonomous system) routing on the internet. Path vector protocol."),
])

        
mock_test(9, [
    ("If today is Wednesday, what day is 100 days later?","100 mod 7 = 2. Wed + 2 = Friday"),
    ("Series: 1, 8, 27, 64, 125, ?","Cubes: 1,2,3,4,5,6=216"),
    ("Logical: A>B, C<D, B=C. Order from smallest to largest?","A>B=C<D. So B=C is middle. A>B means A is larger. D>C means D is larger. Order: B=C < both A and D. Need more info to order A vs D. Answer: C=B < (A and D)  both possible."),
    ("What is the output? x=5; def f(): print(x); x=10; f()","UnboundLocalError  Python sees x=10 inside function, treats x as local, but prints before assignment."),
    ("What does LEFT JOIN return that INNER JOIN does not?","Rows from the left table that have no match in the right table (with NULLs for right table columns)."),
    ("In options, what does positive Theta mean?","Positive Theta means the option gains value with time  unusual; typically sellers have positive Theta (benefit from time decay)."),
    ("A person invests $1000. Loses 50% in year 1, gains 50% in year 2. Final amount?","1000500750. Final=$750 (net 25%)"),
    ("What is a primary key vs unique key?","Primary key: uniquely identifies row, cannot be NULL, only one per table. Unique key: must be unique, CAN be NULL (one), multiple allowed."),
    ("What is BFS time complexity and what data structure does it use?","O(V+E). Uses a Queue."),
    ("Blood relation: Pointing to a woman, a man says 'Her mother is the only daughter of my mother'. Relation?","Only daughter of my mother = my sister. So her mother = my sister. She is my sister's daughter = my niece."),
])

         
mock_test(10, [
    ("EV problem: Roll a die. Win $n if n>4, lose $n if n<=4. EV?","n>4: 525,636. n<=4:11,22,33,44. EV=(25+36)/6+(1234)/6=61/610/6=51/6=$8.50"),
    ("Probability of selecting a committee of 3M and 2W from 6M and 5W?","C(6,3)C(5,2)/C(11,5)=2010/462=200/462=100/231"),
    ("Coding: If FIRE=6935, CODE=?","F=6,I=9,R=3,E=5. C=3,O=15?,D=4,E=5. Pattern: F(6th)=6,I(9th)=9,R(18th)=3(1+8=9no). Try position directly: F=6,I=9,R=183(1+8),E=5(pos). C=3,O=1+5=6,D=4,E=5. CODE=3645"),
    ("A and B run around a 400m track at 4m/s and 6m/s same direction. When do they first meet?","Relative speed=2m/s. Time=400/2=200s"),
    ("What is the Bellman-Ford advantage over Dijkstra?","Handles negative edge weights. Detects negative cycles."),
    ("Python: What is the difference between append and extend for lists?","append adds one element (possibly a list as-is). extend adds all elements of an iterable individually."),
    ("What is arbitrage?","Risk-free profit from simultaneous buy and sell in different markets where price discrepancy exists."),
    ("If np.array has shape (3,4,5), what is the number of elements?","345=60"),
    ("SQL: What is the difference between RANK() and DENSE_RANK()?","RANK() skips ranks after ties (1,1,3). DENSE_RANK() does not skip (1,1,2)."),
    ("Kelly bet: p=0.7 win, b=2 (win 2 per 1 bet). Optimal fraction?","f=(bpq)/b=(20.70.3)/2=(1.40.3)/2=1.1/2=0.55. Bet 55% of bankroll."),
])

page_break()

                                                                                
                                           
                                                                                
h1("Section 15: LAST 2 HOURS  Ultra Revision Sheet")

h2("MENTAL MATH MUST-KNOWS")
bullet("a  b = (a+b)(ab). Always factorise before computing.")
bullet("Successive %: a+b+ab/100. Not just add.")
bullet("Average speed for equal distances: 2xy/(x+y)")
bullet("% of a number: swap  X% of Y = Y% of X")
bullet("Cyclicity: 2,3,7,8 repeat every 4. 4,9 every 2. 0,1,5,6 every 1.")
bullet("Squares to memorise up to 30: 21=441, 22=484, 25=625, 30=900")

h2("PROFIT & LOSS ONE-LINERS")
bullet("SP = CP  (1  P/100). Buy cheap, sell dear.")
bullet("Twin article trap: same SP, +X% and X%  net LOSS = X/100 %")
bullet("Dishonest dealer: Profit = (error)/(statederror)  100")
bullet("Successive discount a,b = a+bab/100")

h2("PROBABILITY MUST-KNOWS")
bullet("P(at least 1) = 1  P(none). Always.")
bullet("Cards: 52 total, 4 aces, 12 face cards, 26 red, 13 per suit.")
bullet("Two dice sum=7 has highest probability (6/36=1/6).")
bullet("P(A and B independent) = P(A)P(B).")

h2("NUMBER THEORY FLASH")
bullet("HCFLCM = product of TWO numbers only.")
bullet("Trailing zeros in n! = number of times 5 divides n! = floor(n/5)+floor(n/25)+...")
bullet("7^1=7,7^2=49,7^3=343,7^4=2401  units: 7,9,3,1 cycle of 4.")
bullet("11 divisibility: alternating digit sum divisible by 11.")

h2("OOP IN 30 SECONDS")
bullet("Overloading=compile time (same name, diff params). Overriding=runtime (child redefines).")
bullet("Abstract class: partial implementation. Interface: pure contract.")
bullet("Java: no multiple class inheritance; yes multiple interface implementation.")
bullet("final class: no subclass. final method: no override. final var: constant.")

h2("DBMS IN 30 SECONDS")
bullet("ACID: Atomicity, Consistency, Isolation, Durability.")
bullet("1NF2NF: remove partial deps. 2NF3NF: remove transitive deps. 3NFBCNF: every determinant is superkey.")
bullet("HAVING filters groups. WHERE filters rows. ORDER BY sorts. GROUP BY aggregates.")
bullet("INNER JOIN: only matches. LEFT JOIN: all left + matches. FULL OUTER: everything.")

h2("OS IN 30 SECONDS")
bullet("Deadlock: Mutual exclusion + Hold&Wait + No preemption + Circular wait  ALL 4 needed.")
bullet("Scheduling: SJF minimises avg waiting. Round Robin is fair. SRTF is preemptive SJF.")
bullet("Page replacement: LRU best practically. FIFO suffers Belady's anomaly.")
bullet("Semaphore: wait() decrements (block if 0). signal() increments (wake waiting).")

h2("ALGO/DS IN 30 SECONDS")
bullet("Binary search: O(log n). Only on sorted array.")
bullet("Quicksort: O(n) worst (sorted input + bad pivot). O(n log n) avg.")
bullet("Merge sort: always O(n log n). Stable. Extra space O(n).")
bullet("HashMap: O(1) average. O(n) worst (all collisions).")
bullet("BFS: Queue. DFS: Stack (or recursion). Both O(V+E).")

h2("PYTHON IN 30 SECONDS")
bullet("list vs tuple: mutable vs immutable. Both ordered.")
bullet("dict.get(key, default): safe access, no KeyError.")
bullet("'is' vs '==': identity vs value. Small ints cached (5 to 256)  'is' may pass for those.")
bullet("List comprehension: [expr for x in iter if cond]. Faster than for+append.")
bullet("*args=tuple, **kwargs=dict. Unpack with *lst and **dct in calls.")

h2("NUMPY IN 30 SECONDS")
bullet("Broadcasting: shapes compatible if equal or one of them is 1.")
bullet("a.reshape(-1): flatten. a.flatten(): returns copy. a.ravel(): returns view.")
bullet("np.axis=0: operate down rows (column-wise). axis=1: across columns (row-wise).")
bullet("a@b or np.dot(a,b): matrix multiplication.")

h2("TRADING/QUANT IN 30 SECONDS")
bullet("EV =  poutcome. Always compute before deciding.")
bullet("Kelly fraction = (bpq)/b. Over-betting is worse than under-betting.")
bullet("Futures: obligation. Options: right not obligation.")
bullet("Long: profit if price rises. Short: profit if price falls.")
bullet("Contango: F>S. Backwardation: F<S. Mark to market: daily settlement.")
bullet("Sharpe = (RRf)/. Higher is better.")
bullet("Vol annualise: _daily  252. _monthly  12.")

h2("LAST 15 MINUTES  Do This:")
body("1. Read every question FULLY before answering  look for qualifiers (at least, exactly, none).")
body("2. Eliminate obviously wrong options first to improve guessing accuracy.")
body("3. For series/pattern: check differences, ratios, squares, alternate terms.")
body("4. For DI: approximate before computing exactly.")
body("5. Check units digit for multiplication answers  eliminates 3 of 4 options often.")
body("6. Blood relation: draw the tree. Don't solve in your head.")
body("7. Time remaining? Skip long calculations, flag and return.")

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(" END OF HANDBOOK  GOOD LUCK! ")
set_font(r, 12, bold=True)

                                                                                
out = "/Users/yashrajshrivastava/Documents/SboxRL/futures_first_handbook.docx"
doc.save(out)
print(f"Saved: {out}")